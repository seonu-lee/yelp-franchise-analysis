"""
Yelp Franchise Analysis — Final Report Word Document Builder
Converts report.md + executive_summary.md + appendix.md into final_report.docx
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

# ─── Paths ───────────────────────────────────────────────────────────────────
BASE    = Path(__file__).parent
FIG_DIR = BASE / "figures"
OUT     = BASE / "final_report.docx"

# ─── Color palette ───────────────────────────────────────────────────────────
C_DARK  = RGBColor(0x1A, 0x1A, 0x2E)   # near-black
C_BLUE  = RGBColor(0x1B, 0x6C, 0xA8)   # heading blue
C_TEAL  = RGBColor(0x15, 0x80, 0x79)   # accent
C_GRAY  = RGBColor(0x55, 0x55, 0x55)   # body gray
C_LIGHT = RGBColor(0xF5, 0xF7, 0xFA)   # table even row
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ─── Helper: set paragraph spacing ───────────────────────────────────────────
def _set_spacing(para, before=0, after=6, line_rule=None, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line_rule and line:
        pf.line_spacing_rule = line_rule
        pf.line_spacing = line

def _set_line15(para):
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para.paragraph_format.line_spacing = Pt(18)   # ~1.5 × 11pt

def _run_font(run, size_pt=11, bold=False, italic=False, color=None):
    run.font.name      = 'Malgun Gothic'
    run.font.size      = Pt(size_pt)
    run.font.bold      = bold
    run.font.italic    = italic
    if color:
        run.font.color.rgb = color
    # Apply East-Asian font too
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

def _set_cell_bg(cell, color_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def _add_page_break(doc):
    para = doc.add_paragraph()
    run  = para.add_run()
    run.add_break(docx_break_type())
    _set_spacing(para, 0, 0)

def docx_break_type():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_BREAK
    return WD_BREAK.PAGE

# ─── Style setup ─────────────────────────────────────────────────────────────
def setup_styles(doc):
    styles = doc.styles

    # Normal
    n = styles['Normal']
    n.font.name = 'Malgun Gothic'
    n.font.size = Pt(11)
    n.font.color.rgb = C_DARK
    n.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    n.paragraph_format.line_spacing = Pt(18)
    n.paragraph_format.space_after  = Pt(6)
    rPr = n.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name  = 'Malgun Gothic'
    h1.font.size  = Pt(16)
    h1.font.bold  = True
    h1.font.color.rgb = C_BLUE
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after  = Pt(8)
    h1.paragraph_format.keep_with_next = True

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name  = 'Malgun Gothic'
    h2.font.size  = Pt(14)
    h2.font.bold  = True
    h2.font.color.rgb = C_TEAL
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after  = Pt(6)
    h2.paragraph_format.keep_with_next = True

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name  = 'Malgun Gothic'
    h3.font.size  = Pt(12)
    h3.font.bold  = True
    h3.font.color.rgb = C_DARK
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after  = Pt(4)
    h3.paragraph_format.keep_with_next = True

    # Caption
    if 'Caption' not in [s.name for s in styles]:
        cap = styles.add_style('Caption', 1)
    else:
        cap = styles['Caption']
    cap.font.name   = 'Malgun Gothic'
    cap.font.size   = Pt(9)
    cap.font.italic = True
    cap.font.color.rgb = C_GRAY
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after  = Pt(10)
    cap.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER

# ─── Cover page ──────────────────────────────────────────────────────────────
def add_cover(doc):
    # Top spacing
    for _ in range(6):
        p = doc.add_paragraph()
        _set_spacing(p, 0, 0)

    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Yelp Franchise Analysis")
    run.font.name  = 'Malgun Gothic'
    run.font.size  = Pt(28)
    run.font.bold  = True
    run.font.color.rgb = C_BLUE
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    _set_spacing(p, 0, 8)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Data-Driven Analysis of Franchise Competition\nand Independent Brand Survival Strategy")
    run.font.name   = 'Malgun Gothic'
    run.font.size   = Pt(14)
    run.font.italic = True
    run.font.color.rgb = C_TEAL
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    _set_spacing(p, 0, 4)

    # Horizontal rule (thick paragraph border)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1B6CA8')
    pBdr.append(bottom)
    pPr.append(pBdr)
    _set_spacing(p, 8, 16)

    # Subtitle KO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("프랜차이즈 경쟁 압력이 독립 브랜드의 평점과 소비자 언어에 미치는 영향")
    run.font.name  = 'Malgun Gothic'
    run.font.size  = Pt(13)
    run.font.bold  = True
    run.font.color.rgb = C_DARK
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    _set_spacing(p, 0, 6)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("거리 가중 경쟁압력 지수(FPI)를 활용한 공간·텍스트 통합 분석")
    r2.font.name   = 'Malgun Gothic'
    r2.font.size   = Pt(11)
    r2.font.italic = True
    r2.font.color.rgb = C_GRAY
    rPr = r2._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    _set_spacing(p2, 0, 24)

    # Metadata block
    meta_items = [
        ("과목",      "마케팅조사"),
        ("조",        "4조"),
        ("분석 지역",  "Las Vegas, NV, USA"),
        ("데이터",    "Yelp Open Dataset"),
        ("작성자",    ""),
        ("제출일",    ""),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_lbl = p.add_run(f"{label}:  ")
        r_lbl.font.name  = 'Malgun Gothic'
        r_lbl.font.size  = Pt(11)
        r_lbl.font.bold  = True
        r_lbl.font.color.rgb = C_BLUE
        r_val = p.add_run(value)
        r_val.font.name  = 'Malgun Gothic'
        r_val.font.size  = Pt(11)
        r_val.font.color.rgb = C_DARK
        for r in (r_lbl, r_val):
            rPr = r._r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        _set_spacing(p, 0, 3)

    # Page break after cover
    doc.add_page_break()


# ─── Page numbers ─────────────────────────────────────────────────────────────
def add_page_numbers(doc):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    section = doc.sections[0]
    footer  = section.footer
    para    = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.extend([fldChar1, instrText, fldChar2])
    run.font.name = 'Malgun Gothic'
    run.font.size = Pt(10)


# ─── Inline markup parser ─────────────────────────────────────────────────────
def _add_inline(para, text: str, base_size=11, base_bold=False, base_color=None):
    """Parse **bold**, *italic*, `code`, and plain text from a single line."""
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    pos = 0
    for m in pattern.finditer(text):
        # plain before
        if m.start() > pos:
            run = para.add_run(text[pos:m.start()])
            _run_font(run, base_size, base_bold, False, base_color)
        if m.group(1).startswith('**'):
            run = para.add_run(m.group(2))
            _run_font(run, base_size, True, False, base_color or C_DARK)
        elif m.group(1).startswith('`'):
            run = para.add_run(m.group(4))
            run.font.name = 'Courier New'
            run.font.size = Pt(base_size - 0.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:  # *italic*
            run = para.add_run(m.group(3))
            _run_font(run, base_size, base_bold, True, base_color)
        pos = m.end()
    if pos < len(text):
        run = para.add_run(text[pos:])
        _run_font(run, base_size, base_bold, False, base_color)


# ─── Table renderer ──────────────────────────────────────────────────────────
def _render_table(doc, rows):
    """rows: list of lists of strings. First row = header."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    # Pad short rows
    rows = [r + [''] * (ncols - len(r)) for r in rows]

    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = tbl.cell(ri, ci)
            cell_text = cell_text.strip().lstrip('|').strip()
            # skip separator rows
            if re.match(r'^[-:| ]+$', cell_text):
                continue

            if ri == 0:
                _set_cell_bg(cell, '1B6CA8')
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(cell_text)
                _run_font(run, 10, True, False, C_WHITE)
            else:
                bg = 'F5F7FA' if ri % 2 == 0 else 'FFFFFF'
                _set_cell_bg(cell, bg)
                para = cell.paragraphs[0]
                _add_inline(para, cell_text, 10, False, C_DARK)

            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)

    _set_spacing(doc.add_paragraph(), 0, 4)


# ─── Figure inserter ─────────────────────────────────────────────────────────
FIGURE_MAP = {
    'fig1_city_comparison':    ("Figure 1", "후보 도시 비교 — 업체 수 및 평균 리뷰 수"),
    'fig2_fpi_formula':        ("Figure 2", "FPI 공식 구성 요소 다이어그램"),
    'fig3_sensitivity':        ("Figure 3", "임계거리 민감도 분석 (PART 1 & PART 2)"),
    'fig4_anova_results':      ("Figure 4", "FPI 구간별 평균 별점 비교 (PART 1 vs PART 2)"),
    'fig5_regression':         ("Figure 5", "OLS 회귀분석 결과 — FPI가 별점에 미치는 영향"),
    'fig6_sentiment_comparison':("Figure 6","VADER 감성 강도 비교 (생존 vs 고전 vs 프랜차이즈)"),
    'fig7_franchise_mention':  ("Figure 7", "수제·독립성 표현 vs 체인 관련 표현 빈도 비교"),
    'fig8_temporal_trend':     ("Figure 8", "시계열 감성 추세: 생존 vs 고전 브랜드 (2005–2017)"),
    'fig9_decline_pathway':    ("Figure 9", "고전 브랜드 하락 경로 재구성"),
    'fig10_plan_vs_actual':    ("Figure 10","분석 계획 대비 실제 수행 결과"),
}

def _insert_figure(doc, key: str):
    """Insert a figure by its base filename key (no extension)."""
    path = FIG_DIR / f"{key}.png"
    if not path.exists():
        return
    fig_num, caption_text = FIGURE_MAP.get(key, ("Figure", key))

    # Container paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, 6, 2)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(5.8))

    # Caption
    cap = doc.add_paragraph(style='Caption')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"{fig_num}. {caption_text}")
    r.font.name  = 'Malgun Gothic'
    r.font.size  = Pt(9)
    r.font.italic = True
    r.font.color.rgb = C_GRAY
    rPr = r._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    _set_spacing(cap, 2, 10)


# ─── Markdown → docx line-by-line renderer ────────────────────────────────────
def render_md(doc, md_text: str, in_appendix=False):
    lines      = md_text.splitlines()
    table_rows = []   # accumulate table rows
    in_code    = False
    code_lines = []
    skip_toc   = False  # skip the markdown TOC block

    def flush_table():
        nonlocal table_rows
        if table_rows:
            # drop separator rows
            cleaned = [r for r in table_rows if not re.match(r'^\s*\|?[-:| ]+\|?\s*$', '|'.join(r))]
            if cleaned:
                _render_table(doc, cleaned)
            table_rows = []

    i = 0
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # ── Skip top YAML / HR lines ────────────────────────────────────────
        if stripped in ('---', '==='):
            flush_table()
            i += 1
            continue

        # ── Skip markdown image lines (we handle figures inline below) ──────
        if stripped.startswith('![') and '](figures/' in stripped:
            flush_table()
            # extract key
            m = re.search(r'\(figures/(.+?)\.png\)', stripped)
            if m:
                _insert_figure(doc, m.group(1))
            i += 1
            continue

        # ── Skip embedded HTML comments ─────────────────────────────────────
        if stripped.startswith('<!--'):
            i += 1
            continue

        # ── Code block ──────────────────────────────────────────────────────
        if stripped.startswith('```'):
            flush_table()
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                # render as indented mono block
                p = doc.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                ind = OxmlElement('w:ind')
                ind.set(qn('w:left'), '360')
                pPr.append(ind)
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F4F4F4')
                pPr.append(shd)
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
                _set_spacing(p, 4, 8)
            i += 1
            continue

        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        # ── Table row ────────────────────────────────────────────────────────
        if stripped.startswith('|'):
            cells = [c.strip() for c in stripped.strip('|').split('|')]
            table_rows.append(cells)
            i += 1
            continue
        else:
            flush_table()

        # ── Empty line ───────────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Horizontal rule ──────────────────────────────────────────────────
        if re.match(r'^[-*_]{3,}$', stripped):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            _set_spacing(p, 6, 10)
            i += 1
            continue

        # ── Headings ─────────────────────────────────────────────────────────
        m4 = re.match(r'^#{4}\s+(.*)', stripped)
        m3 = re.match(r'^#{3}\s+(.*)', stripped)
        m2 = re.match(r'^#{2}\s+(.*)', stripped)
        m1 = re.match(r'^#{1}\s+(.*)', stripped)

        if m4:
            p = doc.add_heading(m4.group(1), level=3)
            p.runs[0].font.size = Pt(11)
            p.runs[0].font.bold = True
            i += 1
            continue
        if m3:
            title = m3.group(1)
            # Skip markdown TOC
            if title.startswith('[') and '](' in title:
                i += 1
                continue
            p = doc.add_heading(title, level=3)
            i += 1
            continue
        if m2:
            title = m2.group(1)
            if title.startswith('[') and '](' in title:
                i += 1
                continue
            p = doc.add_heading(title, level=2)
            i += 1
            continue
        if m1:
            title = m1.group(1)
            if title.startswith('[') and '](' in title:
                i += 1
                continue
            p = doc.add_heading(title, level=1)
            i += 1
            continue

        # ── Blockquote ───────────────────────────────────────────────────────
        if stripped.startswith('>'):
            text = re.sub(r'^>\s*', '', stripped)
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), '480')
            pPr.append(ind)
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '6')
            left.set(qn('w:space'), '4')
            left.set(qn('w:color'), '1B6CA8')
            pBdr.append(left)
            pPr.append(pBdr)
            _add_inline(p, text, 10, False, C_GRAY)
            _set_spacing(p, 2, 2)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = Pt(16)
            i += 1
            continue

        # ── Bullet list ──────────────────────────────────────────────────────
        m_bullet = re.match(r'^[-*+]\s+(.*)', stripped)
        if m_bullet:
            p = doc.add_paragraph(style='List Bullet')
            _add_inline(p, m_bullet.group(1))
            _set_spacing(p, 0, 3)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = Pt(16)
            i += 1
            continue

        # ── Numbered list ────────────────────────────────────────────────────
        m_num = re.match(r'^\d+\.\s+(.*)', stripped)
        if m_num:
            p = doc.add_paragraph(style='List Number')
            _add_inline(p, m_num.group(1))
            _set_spacing(p, 0, 3)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = Pt(16)
            i += 1
            continue

        # ── Math / formula lines ─────────────────────────────────────────────
        if stripped.startswith('$$') or stripped.startswith('\\['):
            text = stripped.replace('$$', '').replace('\\[', '').replace('\\]', '').strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.name  = 'Cambria Math'
            run.font.size  = Pt(11)
            run.font.color.rgb = C_DARK
            _set_spacing(p, 4, 4)
            i += 1
            continue

        # ── Normal paragraph ─────────────────────────────────────────────────
        # Merge continuation lines (lines that don't start with a special char)
        para_lines = [stripped]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].strip()
            if (not nxt or
                nxt.startswith('#') or
                nxt.startswith('|') or
                nxt.startswith('```') or
                nxt.startswith('>') or
                re.match(r'^[-*+]\s', nxt) or
                re.match(r'^\d+\.\s', nxt) or
                nxt.startswith('![') or
                re.match(r'^[-*_]{3,}$', nxt)):
                break
            para_lines.append(nxt)
            j += 1

        full_text = ' '.join(para_lines)
        # Skip pure markdown link lines (TOC)
        if re.match(r'^\d+\.\s+\[.*\]\(#.*\)', full_text):
            i = j
            continue

        p = doc.add_paragraph()
        _add_inline(p, full_text)
        _set_spacing(p, 0, 6)
        _set_line15(p)
        i = j

    flush_table()


# ─── Section: Executive Summary ──────────────────────────────────────────────
def add_executive_summary(doc):
    with open(BASE / 'executive_summary.md', encoding='utf-8') as f:
        md = f.read()
    doc.add_page_break()
    render_md(doc, md)


# ─── Section: Main Report ────────────────────────────────────────────────────
def add_main_report(doc):
    with open(BASE / 'report.md', encoding='utf-8') as f:
        md = f.read()
    # Remove the top title (already on cover) and TOC block
    # Remove first h1 line and everything up to "# 1. Executive Summary"
    lines = md.splitlines()
    start = 0
    for idx, ln in enumerate(lines):
        if re.match(r'^# 1\.', ln.strip()):
            start = idx
            break
    md = '\n'.join(lines[start:])
    doc.add_page_break()
    render_md(doc, md)


# ─── Section: Appendix ───────────────────────────────────────────────────────
def add_appendix(doc):
    with open(BASE / 'appendix.md', encoding='utf-8') as f:
        md = f.read()
    doc.add_page_break()
    render_md(doc, md, in_appendix=True)


# ─── Page setup ──────────────────────────────────────────────────────────────
def setup_page(doc):
    section = doc.sections[0]
    section.page_width   = Cm(21.0)
    section.page_height  = Cm(29.7)
    section.left_margin  = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.top_margin   = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


# ─── Main ─────────────────────────────────────────────────────────────────────
def build():
    print("Building final_report.docx ...")
    doc = Document()
    setup_page(doc)
    setup_styles(doc)

    add_cover(doc)
    add_page_numbers(doc)
    add_executive_summary(doc)
    add_main_report(doc)
    add_appendix(doc)

    doc.save(str(OUT))
    size_kb = OUT.stat().st_size // 1024
    print(f"\nSaved: {OUT}")
    print(f"Size : {size_kb} KB")


if __name__ == '__main__':
    build()
